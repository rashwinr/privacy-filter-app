"""DOCX extractor + redactor."""
from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

from docx import Document  # noqa: E402

from app.extractors.docx import extract_text, redact  # noqa: E402


def _make_docx(path: Path, paragraphs: list[str]) -> Path:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def test_extract_text_concatenates_paragraphs(tmp_path: Path):
    p = _make_docx(tmp_path / "a.docx", ["Hello world", "Second paragraph"])
    out = extract_text(p)
    assert "Hello world" in out
    assert "Second paragraph" in out


def test_extract_then_redact_with_fake_filter(tmp_path: Path, fake_filter):
    body = (
        "Hi, my name is Alice Smith.\n"
        "My email is alice@example.com.\n"
        "Unrelated paragraph with no PII."
    )
    src = _make_docx(tmp_path / "src.docx", body.splitlines())
    out_path = tmp_path / "redacted.docx"

    text = extract_text(src)
    entities = fake_filter.detect(text)
    redact(src, entities, out_path)

    # Re-extract and verify originals are gone.
    redacted_text = extract_text(out_path)
    assert "Alice Smith" not in redacted_text
    assert "alice@example.com" not in redacted_text
    assert "[REDACTED:PRIVATE_PERSON]" in redacted_text
    assert "[REDACTED:PRIVATE_EMAIL]" in redacted_text
    # Untouched paragraph survives intact.
    assert "Unrelated paragraph with no PII" in redacted_text


def test_redact_with_no_entities_preserves_content(tmp_path: Path):
    src = _make_docx(tmp_path / "src.docx", ["Just a normal sentence."])
    out_path = tmp_path / "out.docx"
    redact(src, [], out_path)
    assert "Just a normal sentence." in extract_text(out_path)


def test_redact_handles_table_cells(tmp_path: Path, fake_filter):
    """Table cell paragraphs must also be visited."""
    src = tmp_path / "tbl.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Alice Smith"
    doc.save(str(src))

    text = extract_text(src)
    assert "Alice Smith" in text

    out = tmp_path / "tbl.redacted.docx"
    redact(src, fake_filter.detect(text), out)
    assert "Alice Smith" not in extract_text(out)
