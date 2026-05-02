"""End-to-end API tests using FastAPI's TestClient.

The model singleton is replaced by FakePrivacyFilter via the autouse
fixture in conftest.py, so these tests run in well under a second.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest

pytest.importorskip("fastapi")
pytest.importorskip("httpx")

from fastapi.testclient import TestClient  # noqa: E402

# Importing the app pulls in every extractor — skip the whole module if any
# heavy optional dep is missing on this machine.
pytest.importorskip("pydicom")
pytest.importorskip("docx")
pytest.importorskip("fitz")
pytest.importorskip("PIL")

from app.main import app  # noqa: E402


@pytest.fixture
def client(tmp_data_dir: Path) -> TestClient:
    return TestClient(app)


def test_health_reports_loaded_fake_model(client: TestClient):
    r = client.get("/api/health")
    assert r.status_code == 200
    body = r.json()
    assert body["model_loaded"] is True
    assert body["status"] == "ok"
    assert "fake" in body["model"] or "privacy-filter" in body["model"]


def test_supported_types_lists_known_extensions(client: TestClient):
    r = client.get("/api/supported-types")
    assert r.status_code == 200
    exts = r.json()["extensions"]
    for required in (".txt", ".pdf", ".docx", ".png", ".dcm"):
        assert required in exts


def test_redact_txt_end_to_end(client: TestClient, sample_text: str):
    files = {"file": ("note.txt", sample_text, "text/plain")}
    r = client.post("/api/redact", files=files)
    assert r.status_code == 200, r.text
    body = r.json()

    assert body["filename"] == "note.txt"
    # FakePrivacyFilter rules cover all the entities in sample_text.
    assert body["entity_counts"].get("private_person", 0) >= 1
    assert body["entity_counts"].get("private_email", 0) >= 1
    assert body["entity_counts"].get("private_phone", 0) >= 1

    # The redacted preview must be free of the originals.
    preview = body["text_preview_redacted"]
    assert "Alice Smith" not in preview
    assert "alice@example.com" not in preview
    assert "[REDACTED:" in preview


def test_redact_returns_downloadable_urls(client: TestClient, sample_text: str):
    r = client.post("/api/redact", files={"file": ("note.txt", sample_text, "text/plain")})
    body = r.json()

    orig = client.get(body["original_url"])
    assert orig.status_code == 200
    assert orig.content.decode() == sample_text

    red = client.get(body["redacted_url"])
    assert red.status_code == 200
    assert b"Alice Smith" not in red.content
    assert b"[REDACTED:" in red.content


def test_redact_rejects_unsupported_extension(client: TestClient):
    r = client.post(
        "/api/redact",
        files={"file": ("archive.zip", b"\x00\x01", "application/zip")},
    )
    assert r.status_code == 415
    assert "Unsupported" in r.json()["detail"]


def test_redact_rejects_missing_filename(client: TestClient):
    # Sending no filename in the multipart part triggers FastAPI validation.
    r = client.post(
        "/api/redact",
        files={"file": ("", b"hi", "text/plain")},
    )
    # Either 400 (our handler) or 422 (FastAPI validation) is acceptable.
    assert r.status_code in (400, 422)


def test_download_unknown_kind_404(client: TestClient):
    r = client.get("/api/files/wrongkind/whatever.txt")
    assert r.status_code == 404


def test_download_unknown_key_404(client: TestClient):
    r = client.get("/api/files/uploads/does-not-exist.txt")
    assert r.status_code == 404


def test_root_serves_frontend(client: TestClient):
    r = client.get("/")
    assert r.status_code == 200
    assert "Privacy Filter" in r.text
    assert "/static/app.js" in r.text


def test_redact_docx_roundtrip(client: TestClient, tmp_path: Path):
    from docx import Document

    src = tmp_path / "in.docx"
    doc = Document()
    doc.add_paragraph("Hi, my name is Alice Smith.")
    doc.add_paragraph("Untouched paragraph.")
    doc.save(str(src))

    r = client.post(
        "/api/redact",
        files={"file": (
            "in.docx",
            src.read_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["entity_counts"].get("private_person", 0) >= 1
    assert "Alice Smith" not in (body["text_preview_redacted"] or "")
